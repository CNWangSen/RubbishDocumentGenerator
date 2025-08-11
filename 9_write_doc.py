from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from xpinyin import Pinyin
import os
import pandas as pd
pY = Pinyin()

class Writer:
	def __init__(self):
		self.doc = Document()
		self.StandardFormat()
		self.t_CNT=0
		self.AutoSeq=False
		self.LvCNT=[12,0,0,0,0,0,0,0,0,0]
	def ZeroParagrah(self,style):
		style.font.color.rgb = RGBColor(0, 0, 0)
		style.paragraph_format.line_spacing = Pt(18)#行距
		style.paragraph_format.space_before = Pt(0)
		style.paragraph_format.space_after = Pt(0)
		style.paragraph_format.first_line_indent = Pt(0)
		style.font.italic=False
		style.font.bold=False
	def Font(self,style,Name,Size):
		style._element.rPr.rFonts.set(qn('w:eastAsia'), Name)
		style.font.size = Pt(Size)		
	def StandardFormat(self):
		Normal = self.doc.styles['Normal']
		Normal.font.name = 'Times New Roman'
		
		H1 = self.doc.styles['Heading 1']
		H2 = self.doc.styles['Heading 2']
		H3 = self.doc.styles['Heading 3']
		H4 = self.doc.styles['Heading 4']
		H5 = self.doc.styles['Heading 5']
		H6 = self.doc.styles['Heading 6']
		self.Font(Normal,"宋体",12)#小四
		self.Font(H1,"黑体",14)#四号
		self.Font(H2,"楷体",12)#小四
		self.Font(H3,"宋体",12)
		self.Font(H4,"宋体",12)
		self.Font(H5,"宋体",12)
		self.Font(H6,"宋体",12)
		self.ZeroParagrah(Normal)
		self.ZeroParagrah(H1)
		self.ZeroParagrah(H2)
		self.ZeroParagrah(H3)
		self.ZeroParagrah(H4)
		self.ZeroParagrah(H5)
		self.ZeroParagrah(H6)
		Normal.paragraph_format.first_line_indent = Normal.font.size*2#Cm(0.74)#首行缩进2字符

	def h(self,txt,lv):
		if(self.AutoSeq==False):
			Number=""
			self.LvCNT[lv-1]+=1
			for i in range(lv):
				Number+=str(self.LvCNT[i])+"."
			for i in range(lv,len(self.LvCNT)):
				self.LvCNT[i]=0
			txt=Number+txt
		newp=self.doc.add_heading(txt, level=lv)
	def p(self,txt):
		if(txt==""):
			return
		newp=self.doc.add_paragraph()
		newp.add_run(txt)
	def g(self,figpath,tit=""):
		if(not os.path.exists(figpath)):
			return
		paragraph = self.doc.add_paragraph(style='Caption')  # 使用题注样式
		run = paragraph.add_run()
		run.add_picture(figpath, width=Cm(13))
		if(tit!=""):
			titp = self.doc.add_paragraph(tit)
			titp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	def t(self,ll,tit=""):
		if(tit!=""):
			titp = self.doc.add_paragraph(tit)
			titp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
		tab = self.doc.add_table(rows=len(ll),cols=len(ll[0]),style="Table Grid")
		for row in range(len(ll)):
			cells = tab.rows[row].cells
			for col in range(len(ll[0])):
				cells[col].text = str(ll[row][col])
		self.t_CNT+=1
	def save(self,path):
		self.doc.save(path)

class ToDoc:
	def __init__(self):
		self.W=Writer()
		self.FigCNT=0
		self.TabCNT=0
	def deAI(self,txt):
		return txt.replace("\n","")
	def readTxt(self,path):
		with open(path,"r",encoding="utf-8") as ipt:
			data = ipt.readlines()
		return data
	def xq(self):
		self.W.h("需求分析",2)
		for f in os.listdir("req/modules/"):
			mod = self.readTxt("req/modules/"+f)
			moddesc=mod[0]
			modname = f.split("_")[1].replace(".txt","")
			self.W.h(modname+"需求分析",3)
			if(os.path.exists("text/xq/"+modname+".txt")):
				for context in self.readTxt("text/xq/"+modname+".txt"):
					self.W.p(self.deAI(context))
	def arch(self):
		self.W.h("架构设计",2)
		self.funcarch()
		self.techarch()
	def funcarch(self):
		self.W.h("业务架构",3)
		if(os.path.exists("text/funcarch/main.txt")):
			for context in self.readTxt("text/funcarch/main.txt"):
				self.W.p(self.deAI(context))
	def techarch(self):
		self.W.h("技术架构",3)
		if(os.path.exists("text/techarch/main.txt")):
			for context in self.readTxt("text/techarch/main.txt"):
				self.W.p(self.deAI(context))
	def zb(self):
		self.W.h("指标设计",2)
	def modules(self):
		self.W.h("功能模块设计",2)
		for f in os.listdir("req/modules/"):
			mod = self.readTxt("req/modules/"+f)
			moddesc=mod[0]
			modname = f.split("_")[1].replace(".txt","")
			self.W.h(modname,3)
			self.W.p(self.deAI(moddesc+"该模块框图如下图所示："))
			self.FigCNT+=1
			self.W.g("pic/modules/"+modname+".png",tit="图"+str(self.W.LvCNT[0])+"-"+str(self.FigCNT)+" "+modname+"框图")
			for sub_mod_index in range(1,len(mod)):
				submodname=mod[sub_mod_index].replace("\n","")
				print(submodname)
				self.W.h(submodname,4)
				if(os.path.exists("text/modules/"+modname+"/"+submodname+".txt")):
					for context in self.readTxt("text/modules/"+modname+"/"+submodname+".txt"):
						self.W.p(self.deAI(context))
	def jk(self):
		self.W.h("系统接口",2)
		self.W.h("系统内部接口",3)
		self.W.g("pic/jk/internal.png",tit="图"+str(self.W.LvCNT[0])+"内部接口框图")
		self.jkInEx("text/jk/internal_trans.xlsx",True)
		self.W.h("系统外部接口",3)
		self.W.g("pic/jk/external.png",tit="图"+str(self.W.LvCNT[0])+"外部接口框图")
		self.jkInEx("text/jk/external_trans.xlsx",False)
	def jkInEx(self,path,isInternal=True):
		if(not os.path.exists(path)):
			return
		data = pd.read_excel(path)
		rec=data["接收方"]
		send=data["发送方"]
		para=data["接口参数"]
		func=data["接口功能"]
		para_eng = data["接口参数英文"]

		for i in range(len(rec)):
			s = send[i]
			r = rec[i]
			print(r,s)
			INEX = "EXTERNAL"
			if(isInternal):
				INEX = "INTERNAL"
			uid = "XTKZ-"+INEX+"-"+pY.get_initials(s, '')[0:2]+"T"+pY.get_initials(r, '')[0:2]
			self.W.h(s+"向"+r+"发送数据",4)
			self.W.p("接口标识符："+uid)
			self.W.p("接口特性："+s+"向"+r+"发送数据")
			self.W.p("接口功能："+func[i])
			self.W.p("a)接口优先级：2")
			self.W.p("b)接口类型：发送数据")
			self.W.p("c)接口实体所使用的接口通信方法的特征：")
			self.W.p("1)唯一标识符："+uid)
			self.W.p("2)来源："+s)
			self.W.p("3)接收者："+r)
			self.W.p("4)保密性：高")
			self.W.p("d)接口参数定义：")
			table=[["参数名称","参数类型","参数描述"]]

			p_e_list = para_eng[i].split('-')
			p_list =para[i].split('-')
			for p_i in range(len(p_list)):
				table.append([p_e_list[p_i].replace(" ","_").replace("-","_"),"float",p_list[p_i]])
			self.TabCNT+=1
			self.W.t(table,tit="表"+str(self.W.LvCNT[0])+"-"+str(self.TabCNT)+" "+uid+"接口参数定义表")

	def deploy(self):
		self.W.h("部署方案",2)
		if(os.path.exists("text/deploy/main.txt")):
			for context in self.readTxt("text/deploy/main.txt"):
				self.W.p(self.deAI(context))
	def keytech(self):
		self.W.h("关键技术",2)
		if(os.path.exists("text/keytech/main.txt")):
			for context in self.readTxt("text/keytech/main.txt"):
				self.W.p(self.deAI(context))
	def Gen(self):
		pass
	def Run(self):
		print("xq")
		self.xq()
		print("arch")
		self.arch()
		print("zb")
		self.zb()
		print("modules")
		self.modules()
		print("jk")
		self.jk()
		print("deploy")
		self.deploy()
		print("keytech")
		self.keytech()
		self.W.save("final.docx")

D=ToDoc()
D.Run()